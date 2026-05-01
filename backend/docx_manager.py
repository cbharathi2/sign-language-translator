from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io

_doc = None
_current_text = ""

def initialize_doc():
    global _doc
    _doc = Document()

    # Title
    title = _doc.add_heading("Sign Language to Text Output", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Timestamp
    ts = _doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts.runs[0].font.size = Pt(10)
    ts.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    _doc.add_paragraph("─" * 50)
    _doc.add_heading("Detected Text:", level=2)

def set_text(text: str):
    global _doc, _current_text
    if _doc is None:
        initialize_doc()
    _current_text = text
    save_current_text()

def append_text(text: str):
    global _current_text
    _current_text = text
    save_current_text()

def save_current_text():
    global _doc, _current_text
    if _doc is None:
        initialize_doc()
    # Rebuild document content paragraph
    # Remove old content paragraphs (after heading) and re-add
    # Simple approach: save as file
    try:
        content_para = None
        # Find or add content paragraph (index 4 onwards)
        while len(_doc.paragraphs) > 4:
            p = _doc.paragraphs[-1]._element
            p.getparent().remove(p)
        
        p = _doc.add_paragraph(_current_text)
        p.runs[0].font.size = Pt(12)
        _doc.save("output.docx")
    except Exception as e:
        print(f"Save error: {e}")

def get_current_text() -> str:
    return _current_text

def reset_doc():
    global _doc, _current_text
    _doc = None
    _current_text = ""

def get_doc_bytes() -> bytes:
    global _doc
    if _doc is None:
        initialize_doc()
    buf = io.BytesIO()
    _doc.save(buf)
    buf.seek(0)
    return buf.read()
