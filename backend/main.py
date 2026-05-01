from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import cv2, numpy as np, base64, json, asyncio, os
from datetime import datetime

app = FastAPI(title="Sign Language API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ──────────────── Translation ────────────────
from translator import translate_text

@app.get("/translate")
async def translate(lang: str = Query(...), text: str = Query(...)):
    result = translate_text(text, lang)
    return {"translated": result, "lang": lang}

# ──────────────── DOCX Management ────────────────
from docx_manager import (
    append_text, reset_doc, get_current_text,
    set_text, save_current_text, initialize_doc, get_doc_bytes
)

@app.get("/docx/text")
async def get_text():
    return {"text": get_current_text()}

@app.post("/docx/append")
async def append(body: dict):
    word = body.get("word", "")
    append_text(word)
    return {"text": get_current_text()}

@app.post("/docx/set")
async def set_doc_text(body: dict):
    text = body.get("text", "")
    set_text(text)
    return {"text": get_current_text()}

@app.post("/docx/reset")
async def reset():
    reset_doc()
    return {"text": ""}

@app.get("/docx/download")
async def download():
    save_current_text()
    path = "output.docx"
    if not os.path.exists(path):
        initialize_doc()
        save_current_text()
    return FileResponse(path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=f"sign_language_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

@app.post("/docx/create")
async def create_docx(body: dict):
    """Create DOCX with English text and translation."""
    from docx import Document
    
    english_text = body.get("english_text", "")
    translated_text = body.get("translated_text", "")
    language = body.get("language", "")
    language_name = body.get("language_name", "")
    
    if not english_text or not translated_text:
        raise HTTPException(status_code=400, detail="Missing text fields")
    
    doc = Document()
    doc.add_heading("Sign Language Translation", 0)
    doc.add_paragraph()
    
    doc.add_heading("English (Original)", level=2)
    doc.add_paragraph(english_text)
    doc.add_paragraph()
    
    doc.add_heading(f"{language_name} (Translated)", level=2)
    doc.add_paragraph(translated_text)
    doc.add_paragraph()
    
    doc.add_paragraph(
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        style="Normal"
    )
    
    filename = f"translation_{language}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    
    return FileResponse(
        filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

# ──────────────── WebSocket Sign Detection ────────────────
import mediapipe as mp
from sign_detector import SignDetector

detector = SignDetector()

@app.websocket("/ws/detect")
async def websocket_detect(websocket: WebSocket):
    await websocket.accept()
    try:
        while True:
            data = await websocket.receive_text()
            print(f"Received WebSocket data, length: {len(data)}")
            payload = json.loads(data)
            print(f"Payload keys: {list(payload.keys())}")
            img_data = payload.get("image", "")
            mode = payload.get("mode", "asl")

            if not img_data:
                await websocket.send_text(json.dumps({"letter": "", "confidence": 0}))
                continue

            # Decode base64 image
            if "," in img_data:
                img_data = img_data.split(",")[1]
            img_bytes = base64.b64decode(img_data)
            nparr = np.frombuffer(img_bytes, np.uint8)
            frame = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            print(f"Frame decoded: {frame.shape if frame is not None else 'None'}")

            if frame is None:
                await websocket.send_text(json.dumps({"letter": "", "confidence": 0}))
                continue

            result = await asyncio.to_thread(detector.detect, frame, mode)
            print(f"Detection result: {result}")
            await websocket.send_text(json.dumps(result))

    except WebSocketDisconnect:
        pass
    except Exception as e:
        print(f"WebSocket error: {e}")

@app.get("/health")
async def health():
    return {"status": "ok", "version": "1.0.0"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)